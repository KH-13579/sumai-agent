"""
AI住宅コンシェルジュ 要件定義書 Word生成スクリプト
housing-agent-ontology.jsonld から .docx を生成する
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 色定義 ──────────────────────────────────────────────
C_BLACK   = RGBColor(0x1f, 0x23, 0x28)
C_MUTED   = RGBColor(0x57, 0x60, 0x6a)
C_ACCENT  = RGBColor(0x3b, 0x82, 0xd4)
C_PURPLE  = RGBColor(0x7c, 0x5c, 0xd8)
C_WHITE   = RGBColor(0xff, 0xff, 0xff)
C_SURFACE = RGBColor(0xf7, 0xf8, 0xfa)
C_BORDER  = RGBColor(0xe5, 0xe7, 0xeb)
C_RED_BG  = RGBColor(0xfe, 0xe2, 0xe2)
C_RED_FG  = RGBColor(0xb9, 0x1c, 0x1c)
C_YELLOW  = RGBColor(0xfe, 0xf3, 0xc7)
C_YELLOW_FG = RGBColor(0x92, 0x40, 0x0e)
C_GREEN   = RGBColor(0xd1, 0xfa, 0xe5)
C_GREEN_FG = RGBColor(0x06, 0x5f, 0x46)
C_BLUE_BG = RGBColor(0xdb, 0xea, 0xfe)
C_BLUE_FG = RGBColor(0x1e, 0x40, 0xaf)

PY_EXE = r"C:\Users\AikoKanazawa\AppData\Local\Programs\Python\Python313\python.exe"

# ── ヘルパー ─────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, border_color='E5E7EB'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = C_BLACK
        run.font.bold = True
    return p

def add_para(doc, text, size=10.5, color=None, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color if color else C_BLACK
    if align:
        p.alignment = align
    return p

def add_info_row(doc, label, value, label_color=None, value_color=None):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}:  ')
    r1.font.size = Pt(9.5)
    r1.font.bold = True
    r1.font.color.rgb = label_color if label_color else C_MUTED
    r2 = p.add_run(value)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = value_color if value_color else C_BLACK
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    return p

def add_table_header_row(table, headers, bg=C_ACCENT):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_data_row(table, values, row_idx, shade=False):
    row = table.add_row()
    bg = C_SURFACE if shade else RGBColor(0xff, 0xff, 0xff)
    for i, v in enumerate(values):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        run = p.add_run(str(v))
        run.font.size = Pt(9)

def priority_badge(priority):
    badges = {'Must': '● Must', 'Should': '◎ Should', 'Could': '○ Could'}
    return badges.get(priority, priority)

# ── メイン ───────────────────────────────────────────────
def main():
    with open('housing-agent-ontology.jsonld', encoding='utf-8') as f:
        data = json.load(f)

    graph = data.get('@graph', [])

    # ── ノード分類 ──
    def nodes_of_type(type_suffix):
        result = []
        for node in graph:
            t = node.get('@type', '')
            if isinstance(t, list):
                if any(type_suffix in x for x in t):
                    result.append(node)
            elif type_suffix in t:
                result.append(node)
        return result

    def id_to_title(ref_id):
        for node in graph:
            if node.get('@id') == ref_id:
                return node.get('title', ref_id.split(':')[-1])
        return ref_id.split(':')[-1]

    def refs_to_str(val):
        if not val:
            return '—'
        if isinstance(val, list):
            return ', '.join(id_to_title(v.get('@id', v) if isinstance(v, dict) else v) for v in val)
        if isinstance(val, dict):
            return id_to_title(val.get('@id', ''))
        return str(val)

    # ── Document ──
    doc = Document()

    # ページ余白
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # 既存スタイル調整
    style = doc.styles['Normal']
    style.font.name = 'Meiryo'
    style.font.size = Pt(10.5)

    # ═══════════════════════════════════════════════════
    #  表紙
    # ═══════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI住宅コンシェルジュ')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = C_ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('ソフトウェア要件定義書')
    r2.font.size = Pt(18)
    r2.font.bold = True
    r2.font.color.rgb = C_BLACK

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('初回住宅購入者向け AIエージェント')
    r3.font.size = Pt(13)
    r3.font.color.rgb = C_MUTED

    doc.add_paragraph()
    meta_items = [
        ('ドキュメントバージョン', 'v1.0.0'),
        ('基準日', '2026年6月9日'),
        ('対象ペルソナ', 'A層（初めて購入層）・B層（建替え/住替え層）・C層（土地ありだけ層）'),
        ('テックスタック', 'Azure OpenAI GPT-4系 / LangGraph / Three.js / PLATEAU / e-Gov法令API'),
    ]
    tbl = doc.add_table(rows=len(meta_items), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta_items):
        row = tbl.rows[i]
        set_cell_bg(row.cells[0], C_SURFACE)
        c0 = row.cells[0].paragraphs[0]
        r = c0.add_run(k)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = C_MUTED
        c1 = row.cells[1].paragraphs[0]
        r1 = c1.add_run(v)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = C_BLACK
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  1. 概要
    # ═══════════════════════════════════════════════════
    add_heading(doc, '1. サービス概要', level=1)
    onto = next((n for n in graph if n.get('@id') == 'ha:Ontology'), {})
    add_para(doc, onto.get('description', ''), size=10.5)
    doc.add_paragraph()

    add_heading(doc, '1.1 サービスコンセプト', level=2)
    concepts = [
        ('① 実務レベルのアウトプット', '間取り＋見積＋法規チェックがセット。そのまま商談に持っていける。'),
        ('② 中立性', 'ハウスメーカー非依存、営業バイアスなし。'),
        ('③ マルチAI構成', '間取りAI・見積AI・法規AIをオーケストレーター化。競合が真似しにくい構造。'),
        ('④ UX差別化', '「質問に答えるだけで家ができる」⇒ 設計知識不要。'),
    ]
    for title, desc in concepts:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(title + '  ')
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = C_ACCENT
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)

    add_heading(doc, '1.2 既存サービスとの差別化', level=2)
    diff_headers = ['軸', '既存サービス', '当サービス']
    diff_rows = [
        ['間取り生成', 'あり', '建築基準法連携'],
        ['見積',       '人依存', 'AI生成'],
        ['法規チェック', 'なし', '自動'],
        ['比較', '住宅会社比較', 'プラン比較'],
        ['連絡', '問い合わせ', '自作⇒住宅メーカーへ持込'],
    ]
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    add_table_header_row(tbl2, diff_headers)
    for i, row in enumerate(diff_rows):
        add_table_data_row(tbl2, row, i+1, shade=(i % 2 == 0))
    doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  2. ステークホルダー / ペルソナ
    # ═══════════════════════════════════════════════════
    add_heading(doc, '2. ステークホルダー / ペルソナ', level=1)

    personas = nodes_of_type('Persona')
    stakeholders = [n for n in nodes_of_type('Stakeholder') if 'Persona' not in n.get('@type', '')]

    add_heading(doc, '2.1 ユーザーペルソナ', level=2)
    persona_headers = ['項目', 'A層：初めて購入層', 'B層：建替え・住替え層', 'C層：土地ありだけ層']
    persona_keys = [
        ('ageRange', '年齢'),
        ('annualIncome', '年収'),
        ('familyStructure', '家族構成'),
        ('role', '役割'),
        ('interests', '検討状況'),
    ]
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = 'Table Grid'
    add_table_header_row(tbl3, persona_headers, bg=C_PURPLE)

    for key, label in persona_keys:
        row = tbl3.add_row()
        set_cell_bg(row.cells[0], C_SURFACE)
        r0 = row.cells[0].paragraphs[0].add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9)
        for ci, p_node in enumerate(personas[:3], start=1):
            val = p_node.get(key, '—')
            set_cell_borders(row.cells[ci])
            run = row.cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)

    doc.add_paragraph()
    add_heading(doc, '2.2 ペルソナ別 痛み・期待', level=2)
    for p_node in personas[:3]:
        p_title = doc.add_paragraph()
        rh = p_title.add_run(f"【{p_node.get('title', '')}】")
        rh.font.bold = True
        rh.font.size = Pt(10.5)
        rh.font.color.rgb = C_ACCENT
        pains = p_node.get('painPoints', [])
        expects = p_node.get('expectations', [])
        if pains:
            pa = doc.add_paragraph()
            pa.add_run('主な痛み: ').font.bold = True
            for pain in pains:
                bp = doc.add_paragraph(style='List Bullet')
                bp.add_run(pain).font.size = Pt(9.5)
        if expects:
            pe = doc.add_paragraph()
            pe.add_run('このサービスへの期待: ').font.bold = True
            for exp in expects:
                be = doc.add_paragraph(style='List Bullet')
                be.add_run(exp).font.size = Pt(9.5)
        doc.add_paragraph()

    add_heading(doc, '2.3 その他ステークホルダー', level=2)
    stk_headers = ['ID', '名称', '役割', '関心事']
    tbl4 = doc.add_table(rows=1, cols=4)
    tbl4.style = 'Table Grid'
    add_table_header_row(tbl4, stk_headers)
    for i, s in enumerate(stakeholders):
        add_table_data_row(tbl4, [
            s.get('stakeholderId', '—'),
            s.get('title', '—'),
            s.get('role', '—'),
            s.get('interests', '—'),
        ], i+1, shade=(i % 2 == 0))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  3. ゴール
    # ═══════════════════════════════════════════════════
    add_heading(doc, '3. ビジネス・システムゴール', level=1)
    goals = nodes_of_type('Goal')
    goal_headers = ['ゴールID', 'タイトル', '種別', '説明', '測定指標']
    tbl5 = doc.add_table(rows=1, cols=5)
    tbl5.style = 'Table Grid'
    add_table_header_row(tbl5, goal_headers)
    for i, g in enumerate(goals):
        add_table_data_row(tbl5, [
            g.get('goalId', '—'),
            g.get('title', '—'),
            g.get('goalType', '—'),
            g.get('description', '—'),
            g.get('measuredBy', '—'),
        ], i+1, shade=(i % 2 == 0))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  4. 機能要件
    # ═══════════════════════════════════════════════════
    add_heading(doc, '4. 機能要件 (Functional Requirements)', level=1)
    frs = nodes_of_type('FunctionalRequirement')
    priority_order = {'Must': 0, 'Should': 1, 'Could': 2}
    frs.sort(key=lambda x: priority_order.get(x.get('priority', 'Could'), 3))

    for fr in frs:
        req_id   = fr.get('requirementId', '?')
        title    = fr.get('title', '?')
        priority = fr.get('priority', '?')
        mvp      = fr.get('mvpFlag', False)
        desc     = fr.get('description', '—')
        criteria = fr.get('acceptanceCriteria', '—')
        rationale= fr.get('rationale', '')
        depends  = refs_to_str(fr.get('dependsOn'))
        realizes_val = refs_to_str(fr.get('realizes'))
        verified = refs_to_str(fr.get('verifiedBy'))

        # 見出し行
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(10)
        rh = p_head.add_run(f'{req_id}  {title}')
        rh.font.size = Pt(11)
        rh.font.bold = True
        rh.font.color.rgb = C_BLACK

        # バッジ
        p_badge = doc.add_paragraph()
        p_badge.paragraph_format.space_before = Pt(0)
        p_badge.paragraph_format.space_after = Pt(2)
        badge_colors = {'Must': (C_RED_BG, C_RED_FG), 'Should': (C_YELLOW, C_YELLOW_FG), 'Could': (C_GREEN, C_GREEN_FG)}
        bc, fc = badge_colors.get(priority, (C_SURFACE, C_MUTED))
        rb = p_badge.add_run(f' {priority_badge(priority)} ')
        rb.font.size = Pt(9)
        rb.font.bold = True
        rb.font.color.rgb = fc
        if mvp:
            rb2 = p_badge.add_run('  ★ MVP')
            rb2.font.size = Pt(9)
            rb2.font.bold = True
            rb2.font.color.rgb = C_BLUE_FG

        tbl_fr = doc.add_table(rows=2, cols=2)
        tbl_fr.style = 'Table Grid'

        labels = ['説明', '受入基準', '実装理由', '依存関係', '実現ゴール', '検証テスト']
        values = [desc, criteria, rationale or '—', depends, realizes_val, verified]
        pairs = list(zip(labels, values))

        for ri in range(3):
            row = tbl_fr.rows[ri] if ri < len(tbl_fr.rows) else tbl_fr.add_row()
            for ci, (lbl, val) in enumerate(pairs[ri*2:ri*2+2]):
                cell = row.cells[ci]
                set_cell_bg(cell, C_SURFACE if ci == 0 else RGBColor(0xff, 0xff, 0xff))
                set_cell_borders(cell)
                p = cell.paragraphs[0]
                rl = p.add_run(lbl + ': ')
                rl.font.bold = True
                rl.font.size = Pt(8.5)
                rl.font.color.rgb = C_MUTED
                rv = p.add_run(val)
                rv.font.size = Pt(9)

        # 3行目追加
        row3 = tbl_fr.add_row()
        for ci, (lbl, val) in enumerate(pairs[4:6]):
            cell = row3.cells[ci]
            set_cell_bg(cell, C_SURFACE if ci == 0 else RGBColor(0xff, 0xff, 0xff))
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            rl = p.add_run(lbl + ': ')
            rl.font.bold = True
            rl.font.size = Pt(8.5)
            rl.font.color.rgb = C_MUTED
            rv = p.add_run(val)
            rv.font.size = Pt(9)

        doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  5. 非機能要件
    # ═══════════════════════════════════════════════════
    add_heading(doc, '5. 非機能要件 / セキュリティ要件 / 性能要件', level=1)

    nfrs = nodes_of_type('NonFunctionalRequirement')
    srs  = nodes_of_type('SecurityRequirement')
    prs  = nodes_of_type('PerformanceRequirement')

    sections_nfp = [
        ('5.1 非機能要件 (NFR)', nfrs, C_PURPLE),
        ('5.2 セキュリティ要件 (SR)', srs, C_RED_FG),
        ('5.3 性能要件 (PR)', prs, RGBColor(0x85, 0x4d, 0x0e)),
    ]
    for sec_title, nodes, color in sections_nfp:
        add_heading(doc, sec_title, level=2)
        if not nodes:
            doc.add_paragraph('（該当なし）')
            continue
        headers = ['ID', 'タイトル', '優先度', '説明 / 受入基準']
        tbl_n = doc.add_table(rows=1, cols=4)
        tbl_n.style = 'Table Grid'
        add_table_header_row(tbl_n, headers, bg=color)
        for i, n in enumerate(nodes):
            desc_full = n.get('description', '—')
            criteria  = n.get('acceptanceCriteria', '')
            combined  = desc_full + ('\n受入基準: ' + criteria if criteria else '')
            row = tbl_n.add_row()
            bg = C_SURFACE if i % 2 == 0 else RGBColor(0xff, 0xff, 0xff)
            for ci, val in enumerate([
                n.get('requirementId', '—'),
                n.get('title', '—'),
                priority_badge(n.get('priority', '—')),
                combined,
            ]):
                cell = row.cells[ci]
                set_cell_bg(cell, bg)
                set_cell_borders(cell)
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(9)
        doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  6. ユースケース
    # ═══════════════════════════════════════════════════
    add_heading(doc, '6. ユースケース', level=1)
    usecases = nodes_of_type('UseCase')

    for uc in usecases:
        uc_id    = uc.get('useCaseId', '?')
        uc_title = uc.get('title', '?')
        add_heading(doc, f'{uc_id}: {uc_title}', level=2)

        infos = [
            ('説明', uc.get('description', '—')),
            ('アクター', refs_to_str(uc.get('actor'))),
            ('トリガー', uc.get('trigger', '—')),
            ('事前条件', uc.get('precondition', '—')),
            ('事後条件', uc.get('postcondition', '—')),
        ]
        for lbl, val in infos:
            add_info_row(doc, lbl, val)

        main_flow = uc.get('mainFlow', [])
        if main_flow:
            p_mf = doc.add_paragraph()
            p_mf.add_run('メインフロー:').font.bold = True
            for step in main_flow:
                bp = doc.add_paragraph(style='List Number')
                bp.add_run(step).font.size = Pt(9.5)

        alt_flow = uc.get('alternativeFlow', [])
        if alt_flow:
            p_af = doc.add_paragraph()
            p_af.add_run('代替フロー:').font.bold = True
            for step in alt_flow:
                bp = doc.add_paragraph(style='List Bullet')
                bp.add_run(step).font.size = Pt(9.5)

        includes_val = refs_to_str(uc.get('includes'))
        if includes_val != '—':
            add_info_row(doc, 'include関係', includes_val)
        doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  7. テストケース
    # ═══════════════════════════════════════════════════
    add_heading(doc, '7. テストケース', level=1)
    testcases = nodes_of_type('TestCase')

    tc_headers = ['ID', 'タイトル', '種別', '入力条件', '期待結果', '検証対象要件']
    tbl_tc = doc.add_table(rows=1, cols=6)
    tbl_tc.style = 'Table Grid'
    add_table_header_row(tbl_tc, tc_headers, bg=RGBColor(0x0c, 0x4a, 0x6e))
    for i, tc in enumerate(testcases):
        add_table_data_row(tbl_tc, [
            tc.get('testCaseId', '—'),
            tc.get('title', '—'),
            tc.get('testType', '—'),
            tc.get('testInput', '—'),
            tc.get('expectedOutput', '—'),
            refs_to_str(tc.get('verifies')),
        ], i+1, shade=(i % 2 == 0))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  8. 外部連携システム
    # ═══════════════════════════════════════════════════
    add_heading(doc, '8. 外部連携システム', level=1)
    ext_nodes = nodes_of_type('ExternalSystem')

    ext_headers = ['ID', 'サービス名', 'データ種別', '利用可否', 'API / URL']
    tbl_ext = doc.add_table(rows=1, cols=5)
    tbl_ext.style = 'Table Grid'
    add_table_header_row(tbl_ext, ext_headers, bg=RGBColor(0x37, 0x41, 0x51))
    for i, e in enumerate(ext_nodes):
        add_table_data_row(tbl_ext, [
            e.get('externalSystemId', '—'),
            e.get('title', '—'),
            e.get('dataType', '—'),
            e.get('availability', '—'),
            e.get('apiEndpoint', e.get('description', '—')),
        ], i+1, shade=(i % 2 == 0))
    doc.add_paragraph()
    add_heading(doc, '推奨テックスタック', level=2)
    stack_items = [
        'オーケストレーター: Azure OpenAI (GPT-4系) + LangGraph',
        '間取り自前ソルバー + Three.js (将来: Autodesk Forma)',
        '部材: 建材カタログ PDF の RAG (将来は建材ベンダー API 契約)',
        '見積: 建設物価統計 CSV + 建築各工統計',
        '法規: 国土数値情報 + IOBA 法令 DB の RAG',
        '地図: 国土地理院タイル + PLATEAU',
    ]
    for item in stack_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(item).font.size = Pt(10)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  9. トレーサビリティマトリクス
    # ═══════════════════════════════════════════════════
    add_heading(doc, '9. トレーサビリティマトリクス', level=1)
    add_para(doc, '各要件とビジネスゴールの対応関係、MVPフラグを示す。', size=9.5, color=C_MUTED)
    doc.add_paragraph()

    tr_headers = ['要件ID', 'タイトル', 'GOAL-001\n中立支援', 'GOAL-002\n一気通貫', 'GOAL-003\n差別化', 'MVP', '優先度']
    tbl_tr = doc.add_table(rows=1, cols=7)
    tbl_tr.style = 'Table Grid'
    add_table_header_row(tbl_tr, tr_headers)

    all_reqs = (
        nodes_of_type('FunctionalRequirement') +
        nodes_of_type('NonFunctionalRequirement') +
        nodes_of_type('SecurityRequirement') +
        nodes_of_type('PerformanceRequirement')
    )
    goal_map = {
        'ha:Goal-NeutralConsulting': 'GOAL-001',
        'ha:Goal-OneStopService':    'GOAL-002',
        'ha:Goal-MarketGap':         'GOAL-003',
    }
    for i, req in enumerate(all_reqs):
        realizes_raw = req.get('realizes', [])
        if isinstance(realizes_raw, dict):
            realizes_raw = [realizes_raw]
        realized_ids = set()
        for r in realizes_raw:
            if isinstance(r, dict):
                realized_ids.add(r.get('@id', ''))
            else:
                realized_ids.add(r)
        g1 = '✓' if 'ha:Goal-NeutralConsulting' in realized_ids else ''
        g2 = '✓' if 'ha:Goal-OneStopService'    in realized_ids else ''
        g3 = '✓' if 'ha:Goal-MarketGap'         in realized_ids else ''
        mvp_flag = '●' if req.get('mvpFlag', False) else ''
        row = tbl_tr.add_row()
        vals = [
            req.get('requirementId', '—'),
            req.get('title', '—'),
            g1, g2, g3,
            mvp_flag,
            req.get('priority', '—'),
        ]
        shade = i % 2 == 0
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            set_cell_bg(cell, C_SURFACE if shade else RGBColor(0xff, 0xff, 0xff))
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if ci in (2, 3, 4, 5):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.color.rgb = C_ACCENT if val == '✓' else (C_BLUE_FG if val == '●' else C_BLACK)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    #  10. オントロジー設計メモ (JSON-LD スニペット)
    # ═══════════════════════════════════════════════════
    add_heading(doc, '10. オントロジー設計メモ', level=1)
    add_heading(doc, '10.1 クラス階層', level=2)
    classes = [
        ('sro:Requirement', '要件の基底クラス'),
        ('  sro:FunctionalRequirement', 'サブクラス: 機能要件'),
        ('  sro:NonFunctionalRequirement', 'サブクラス: 非機能要件'),
        ('  sro:SecurityRequirement', 'サブクラス: セキュリティ要件'),
        ('  sro:PerformanceRequirement', 'サブクラス: 性能要件'),
        ('sro:Stakeholder', '利害関係者'),
        ('  ha:Persona', 'サブクラス: ユーザーペルソナ'),
        ('sro:Goal', 'ビジネス/システム目標'),
        ('sro:UseCase', 'ユースケース'),
        ('sro:TestCase', 'テストケース'),
        ('sro:ExternalSystem', '外部連携システム'),
    ]
    for cls, desc in classes:
        p = doc.add_paragraph()
        r1 = p.add_run(cls.ljust(40))
        r1.font.name = 'Courier New'
        r1.font.size = Pt(9)
        r1.font.color.rgb = C_PURPLE
        r2 = p.add_run('  ' + desc)
        r2.font.size = Pt(9)
        r2.font.color.rgb = C_MUTED
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

    add_heading(doc, '10.2 要件間リレーションシッププロパティ', level=2)
    rels = [
        ('sro:dependsOn',    '依存関係 — この要件はX無しには実現できない'),
        ('sro:conflictsWith','競合関係 — 同時実装に設計上の緊張が生じる'),
        ('sro:refines',      '精緻化 — 上位要件を具体化する'),
        ('sro:realizes',     '実現 — この要件がゴールを達成する'),
        ('sro:verifiedBy',   '検証 — この要件を検証するテストケース'),
        ('sro:derivedFrom',  '導出元 — ペルソナ課題やゴールからの派生'),
        ('sro:includes',     'UC include関係'),
        ('sro:extends',      'UC extend関係'),
    ]
    rel_headers = ['プロパティ', '意味']
    tbl_rel = doc.add_table(rows=1, cols=2)
    tbl_rel.style = 'Table Grid'
    add_table_header_row(tbl_rel, rel_headers)
    for i, (prop, meaning) in enumerate(rels):
        add_table_data_row(tbl_rel, [prop, meaning], i+1, shade=(i%2==0))

    doc.add_paragraph()

    # ── フッター ──
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = p_footer.add_run('Generated by IBM Bob  |  AI住宅コンシェルジュ 要件定義書 v1.0.0  |  2026-06-09')
    rf.font.size = Pt(8)
    rf.font.color.rgb = C_MUTED

    # ── 保存 ──
    out_path = 'AI住宅コンシェルジュ_要件定義書_v1.0.0.docx'
    doc.save(out_path)
    print(f'[OK] Generated: {out_path}')

if __name__ == '__main__':
    main()
